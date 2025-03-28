import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from googletrans import Translator
import io
import base64
from datetime import datetime
import json
import re
import random
import time

# For the Case Finder tab (add near the top with other imports)
try:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.chrome.service import Service
    from webdriver_manager.chrome import ChromeDriverManager
    from selenium.webdriver.chrome.options import Options
    import urllib.parse
    SELENIUM_AVAILABLE = True
except ImportError:
    SELENIUM_AVAILABLE = False

# Load environment variables
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")
if not api_key:
    st.error("GOOGLE_API_KEY is not set. Please check your .env file.")
    st.stop()

genai.configure(api_key=api_key)
translator = Translator()

# Custom CSS
def load_css():
    st.markdown("""
    <style>
    /* Modern Legal App Styling */
    
    /* Base styling */
    @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;500;700&display=swap');
    @import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;700&display=swap');
    
    * {
        font-family: 'Roboto', sans-serif;
        transition: all 0.3s ease;
    }
    
    /* Main container styling with gradient */
    .main {
        background: linear-gradient(135deg, #f5f7fa 0%, #e4e8f0 100%);
        padding: 30px;
        color: #333;
        min-height: 100vh;
    }
    
    /* Header styling with shadow */
    .stApp header {
        background: linear-gradient(90deg, #0a2540 0%, #1a3c5a 100%) !important;
        box-shadow: 0 4px 12px rgba(0,0,0,0.2) !important;
    }
    
    h1, h2, h3 {
        font-family: 'Playfair Display', serif;
        color: #0a2540;
    }
    
    /* App title with elegant animation */
    .app-title {
        text-align: center;
        background: linear-gradient(90deg, #0a2540, #2a5a85, #0a2540);
        background-size: 200% auto;
        color: #fff;
        padding: 20px;
        border-radius: 8px;
        margin-bottom: 30px;
        box-shadow: 0 6px 15px rgba(10, 37, 64, 0.2);
        animation: gradient 8s ease infinite;
        display: flex;
        align-items: center;
        justify-content: center;
    }
    
    @keyframes gradient {
        0% {background-position: 0% 50%;}
        50% {background-position: 100% 50%;}
        100% {background-position: 0% 50%;}
    }
    
    .app-title h1 {
        color: white;
        margin: 0;
        font-size: 2.5rem;
    }
    
    .app-icon {
        font-size: 2.8rem;
        margin-right: 15px;
    }
    
    /* Custom animated card for chat messages */
    .chat-message {
        padding: 1.5rem;
        border-radius: 0.8rem;
        margin-bottom: 1.5rem;
        display: flex;
        flex-direction: column;
        animation: slideIn 0.3s ease-out;
        transition: transform 0.2s;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
        width: 100%;
        overflow-wrap: break-word;
        word-wrap: break-word;
        word-break: break-word;
        color: #333; /* Add explicit text color */
    }
    
    .message-content {
        width: 100%;
        padding: 0.5rem;
        color: #333; /* Add explicit text color */
    }
    
    @keyframes slideIn {
        from {opacity: 0; transform: translateY(20px);}
        to {opacity: 1; transform: translateY(0);}
    }
    
    .chat-message:hover {
        transform: translateY(-5px);
        box-shadow: 0 6px 12px rgba(0,0,0,0.15);
    }
    
    .chat-message.user {
       
        border-left: 5px solid #2e86de;
        color: #000000; /* Add explicit text color */
    }
    
    .chat-message.bot {
        
        border-left: 5px solid #5cb85c;
        color: #000000; /* Add explicit text color */
    }
    
    /* Elegant sidebar styling with depth */
    .css-1d391kg, .css-hxt7ib {
        background: linear-gradient(180deg, #f1f3f8 0%, #e4e8f0 100%);
        box-shadow: inset -5px 0 10px -5px rgba(0,0,0,0.1);
    }
    
    /* Modern File uploader */
    .stFileUploader {
        padding: 20px;
        border: 2px dashed #1a3c5a;
        border-radius: 15px;
        text-align: center;
        background-color: rgba(26, 60, 90, 0.05);
        transition: all 0.3s;
    }
    
    .stFileUploader:hover {
        background-color: rgba(26, 60, 90, 0.1);
        border-color: #2a5a85;
    }
    
    /* Animated Buttons */
    .stButton>button {
        background: linear-gradient(90deg, #1a3c5a 0%, #2a5a85 100%);
        color: white;
        font-weight: 500;
        border: none;
        border-radius: 50px;
        padding: 0.7rem 1.5rem;
        transition: all 0.3s;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    
    .stButton>button:hover {
        background: linear-gradient(90deg, #2a5a85 0%, #3a7ab5 100%);
        box-shadow: 0 6px 12px rgba(0,0,0,0.2);
        transform: translateY(-2px);
    }
    
    .stButton>button:active {
        transform: translateY(1px);
        box-shadow: 0 2px 4px rgba(0,0,0,0.2);
    }
    
    /* History section with scrolling effect */
    .history-container {
        max-height: 600px;
        overflow-y: auto;
        padding: 15px;
        background: linear-gradient(to bottom, #fff, #f8f9fa);
        border-radius: 8px;
        border: 1px solid #e9ecef;
        box-shadow: inset 0 0 10px rgba(0,0,0,0.05);
        scrollbar-width: thin;
    }
    
    .history-container::-webkit-scrollbar {
        width: 6px;
    }
    
    .history-container::-webkit-scrollbar-track {
        background: #f1f1f1;
        border-radius: 10px;
    }
    
    .history-container::-webkit-scrollbar-thumb {
        background: #1a3c5a;
        border-radius: 10px;
    }
    
    .history-item {
        margin-bottom: 12px;
        padding: 20px;
        background-color: white;
        border-radius: 8px;
        border-left: 4px solid #1a3c5a;
        box-shadow: 0 2px 8px rgba(0,0,0,0.06);
        transition: all 0.2s;
        animation: fadeIn 0.5s ease-out;
        color: #ffffff;
    }
    
    .history-question {
        font-weight: 500;
        color: #1a3c5a;
        margin-bottom: 12px;
        font-size: 1.1em;
    }
    
    .history-answer {
        color: #333;
        background-color: #f8f9fa;
        padding: 10px;
        border-radius: 4px;
        border-left: 3px solid #5cb85c;
        margin-top: 8px;
    }
    
    @keyframes fadeIn {
        from {opacity: 0;}
        to {opacity: 1;}
    }
    
    .history-item:hover {
        transform: scale(1.01);
        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        background-color: #f9fbfd;
    }
    
    .history-timestamp {
        font-size: 0.8rem;
        color: #6c757d;
        margin-bottom: 5px;
    }
    
    .history-question {
        font-weight: 500;
        color: #1a3c5a;
        margin-bottom: 8px;
    }
    
    .history-answer {
        color: #495057;
    }
    
    /* Document tab with glass morphism effect */
    .document-tab {
        padding: 25px;
        background-color: rgba(255, 255, 255, 0.8);
        backdrop-filter: blur(10px);
        border-radius: 12px;
        box-shadow: 0 8px 32px rgba(31, 38, 135, 0.1);
        border: 1px solid rgba(255, 255, 255, 0.18);
        margin-bottom: 30px;
        transition: all 0.3s;
    }
    
    .document-tab:hover {
        box-shadow: 0 8px 32px rgba(31, 38, 135, 0.2);
        transform: translateY(-5px);
    }
    
    /* Contract template with elegant styling */
    .contract-template {
        padding: 30px;
        background-color: #f8f8f8;
        background-image: linear-gradient(45deg, #f9f9f9 25%, transparent 25%, transparent 75%, #f9f9f9 75%, #f9f9f9), 
                        linear-gradient(45deg, #f9f9f9 25%, transparent 25%, transparent 75%, #f9f9f9 75%, #f9f9f9);
        background-size: 20px 20px;
        background-position: 0 0, 10px 10px;
        border: 1px solid #ddd;
        border-radius: 8px;
        font-family: 'Playfair Display', serif;
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
    }
    
    /* Timeline styling with 3D effect */
    .timeline-container {
        position: relative;
        margin: 40px 0;
        perspective: 1000px;
    }
    
    .timeline-container::after {
        content: '';
        position: absolute;
        width: 6px;
        background: linear-gradient(to bottom, #1a3c5a, #2a5a85);
        top: 0;
        bottom: 0;
        left: 50%;
        margin-left: -3px;
        border-radius: 3px;
        box-shadow: 0 0 10px rgba(26, 60, 90, 0.3);
    }
    
    .timeline-item {
        padding: 15px 50px;
        position: relative;
        background-color: inherit;
        width: 50%;
        perspective: 1000px;
        transform-style: preserve-3d;
        animation: timelineAppear 0.8s ease-out;
    }
    
    @keyframes timelineAppear {
        from {opacity: 0; transform: rotateY(-30deg) translateX(-50px);}
        to {opacity: 1; transform: rotateY(0) translateX(0);}
    }
    
    .timeline-item::after {
        content: '';
        position: absolute;
        width: 25px;
        height: 25px;
        right: -17px;
        background: linear-gradient(135deg, #ffffff 0%, #f1f3f8 100%);
        border: 4px solid #1a3c5a;
        top: 15px;
        border-radius: 50%;
        z-index: 1;
        box-shadow: 0 0 0 5px rgba(26, 60, 90, 0.2);
        transition: all 0.3s;
    }
    
    .timeline-item:hover::after {
        box-shadow: 0 0 0 8px rgba(26, 60, 90, 0.3);
        transform: scale(1.2);
    }
    
    .timeline-item.left {
        left: 0;
        animation-delay: 0.2s;
    }
    
    .timeline-item.right {
        left: 50%;
        animation-delay: 0.4s;
    }
    
    .timeline-item.right::after {
        left: -16px;
    }
    
    .timeline-content {
        padding: 20px 30px;
        background: linear-gradient(to bottom right, #ffffff, #f5f7fa);
        position: relative;
        border-radius: 12px;
        box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        transition: all 0.3s;
        transform: translateZ(0);
    }
    
    .timeline-content:hover {
        box-shadow: 0 8px 25px rgba(0,0,0,0.15);
        transform: translateY(-5px) translateZ(10px);
    }
    
    .timeline-date {
        font-weight: 700;
        color: #1a3c5a;
        margin-bottom: 10px;
        font-size: 1.1rem;
        border-bottom: 2px solid #1a3c5a;
        padding-bottom: 5px;
        display: inline-block;
    }
    
    /* Animated Download button */
    .download-button {
        display: inline-block;
        padding: 12px 24px;
        background: linear-gradient(90deg, #1a3c5a 0%, #2a5a85 100%);
        color: white;
        text-decoration: none;
        border-radius: 50px;
        font-weight: 500;
        margin-top: 15px;
        text-align: center;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
        transition: all 0.3s;
        position: relative;
        overflow: hidden;
    }
    
    .download-button::before {
        content: "";
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(255,255,255,0.4), transparent);
        transition: all 0.6s;
    }
    
    .download-button:hover {
        background: linear-gradient(90deg, #2a5a85 0%, #3a7ab5 100%);
        box-shadow: 0 6px 12px rgba(0,0,0,0.2);
        transform: translateY(-3px);
        color: white;
        text-decoration: none;
    }
    
    .download-button:hover::before {
        left: 100%;
    }
    
    .download-button:active {
        transform: translateY(1px);
        box-shadow: 0 2px 4px rgba(0,0,0,0.2);
    }
    
    /* Tab navigation styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 10px;
        background-color: rgba(255, 255, 255, 0.7);
        backdrop-filter: blur(10px);
        padding: 10px;
        border-radius: 50px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
    }

    .stTabs [data-baseweb="tab"] {
        height: 50px;
        border-radius: 25px;
        margin-right: 5px;
        background-color: rgba(255, 255, 255, 0.5);
        color: #1a3c5a;
        transition: all 0.3s;
        font-weight: 500;
        border: 1px solid transparent;
        display: flex;
        align-items: center;
        justify-content: center;
    }

    .stTabs [aria-selected="true"] {
        background-color: #1a3c5a !important;
        color: white !important;
        border: 1px solid #1a3c5a;
        box-shadow: 0 4px 10px rgba(26, 60, 90, 0.3);
        transform: translateY(-2px);
    }

    .stTabs [data-baseweb="tab"]:hover {
        background-color: rgba(26, 60, 90, 0.1);
        transform: translateY(-2px);
        box-shadow: 0 4px 10px rgba(0,0,0,0.1);
    }
    
    /* Animated text inputs */
    .stTextInput>div>div>input,
    .stTextArea>div>div>textarea, 
    .stSelectbox>div>div>div {
        border-radius: 8px;
        border: 2px solid #e1e5eb;
        padding: 12px 15px;
        transition: all 0.3s;
        background-color: rgba(255, 255, 255, 0.8);
        box-shadow: 0 2px 5px rgba(0,0,0,0.05);
    }

    .stTextInput>div>div>input:focus,
    .stTextArea>div>div>textarea:focus,
    .stSelectbox>div>div>div:focus {
        border-color: #1a3c5a;
        box-shadow: 0 0 0 2px rgba(26, 60, 90, 0.2);
        transform: translateY(-2px);
    }

    .stTextInput>div>div>input:hover,
    .stTextArea>div>div>textarea:hover,
    .stSelectbox>div>div>div:hover {
        border-color: #2a5a85;
    }
    
    /* Improved selectbox styling for dynamic height */
    .stSelectbox>div>div>div {
        border-radius: 4px;
        border: 1.8px solid #e1e5eb;
        padding: 12px 15px;
        transition: all 0.3s;
        background-color: rgba(255, 255, 255, 0.8);
        box-shadow: 0 2px 5px rgba(0,0,0,0.05);
        min-height: 60px;
        height: auto;
        display: flex;
        align-items: center;
        line-height: 1.4;
    }

    /* Style for dropdown options to ensure consistent text alignment */
    .stSelectbox div[data-baseweb="select"] ul li {
        padding: 10px 15px;
        line-height: 1.4;
        min-height: 72px;
        height: auto;
        display: flex;
        align-items: center;
    }

    /* Fix for dropdown arrow alignment */
    .stSelectbox div[data-baseweb="select"] div[role="button"] {
        display: flex;
        align-items: center;
    }
    
    /* Progress bar styling */
    .stProgress > div > div > div > div {
        background-color: #1a3c5a;
        background-image: linear-gradient(45deg, rgba(255, 255, 255, 0.15) 25%, transparent 25%, transparent 50%, rgba(255, 255, 255, 0.15) 50%, rgba(255, 255, 255, 0.15) 75%, transparent 75%, transparent);
        background-size: 1rem 1rem;
        animation: progress-bar-stripes 1s linear infinite;
    }
    
    @keyframes progress-bar-stripes {
        0% {background-position-x: 1rem;}
    }

    /* Loading spinner with animation */
    .stSpinner {
        border: 4px solid rgba(26, 60, 90, 0.1);
        border-radius: 50%;
        border-top: 4px solid #1a3c5a;
        width: 40px;
        height: 40px;
        animation: spinner 1s linear infinite;
        margin: 20px auto;
    }
    
    @keyframes spinner {
        0% {transform: rotate(0deg);}
        100% {transform: rotate(360deg);}
    }
    
    /* Custom checkboxes and radio buttons */
    .stCheckbox > div[role="checkbox"],
    input[type="radio"] {
        cursor: pointer;
        width: 20px;
        height: 20px;
        margin-right: 10px;
        position: relative;
        transition: all 0.3s;
    }
    
    .stCheckbox > div[role="checkbox"]:checked,
    input[type="radio"]:checked {
        background-color: #1a3c5a;
        border-color: #1a3c5a;
    }
    
    /* Dark mode adjustments */
    @media (prefers-color-scheme: dark) {
        .main {
            background: linear-gradient(135deg, #1c2331 0%, #121a24 100%);
            color: #e1e5eb;
        }
        
        h1, h2, h3 {
            color: #e1e5eb;
        }
        
        .document-tab {
            background-color: rgba(30, 41, 59, 0.8);
            border: 1px solid rgba(50, 60, 80, 0.5);
        }
        
        .timeline-content {
            background: linear-gradient(to bottom right, #1c2331, #121a24);
            color: #e1e5eb;
        }
        
        .history-container {
            background: linear-gradient(to bottom, #1c2331, #121a24);
            border-color: #2a3441;
        }
        
        .history-item {
            background-color: #1e293b;
            color: #e1e5eb;
        }
        
        .stTextInput>div>div>input,
        .stTextArea>div>div>textarea, 
        .stSelectbox>div>div>div {
            background-color: rgba(30, 41, 59, 0.8);
            border-color: #2a3441;
            color: #e1e5eb;
        }
        
        .chat-message, .message-content, .chat-message.user, .chat-message.bot {
            color: #e1e5eb; /* Light color for text in dark mode */
        }
    }
    
    /* UI animations and interactions for legal document elements */
    .contract-section {
        margin-bottom: 20px;
        padding: 15px;
        border-radius: 8px;
        background-color: #fff;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        position: relative;
        overflow: hidden;
        transition: all 0.3s;
    }
    
    .contract-section::before {
        content: "";
        position: absolute;
        top: 0;
        left: 0;
        height: 100%;
        width: 5px;
        background-color: #1a3c5a;
        transform: scaleY(0);
        transition: transform 0.3s ease;
    }
    
    .contract-section:hover {
        transform: translateX(5px);
        box-shadow: 0 5px 15px rgba(0,0,0,0.15);
    }
    
    .contract-section:hover::before {
        transform: scaleY(1);
    }
    
    .contract-heading {
        font-family: 'Playfair Display', serif;
        color: #1a3c5a;
        border-bottom: 2px solid #e9ecef;
        padding-bottom: 10px;
        margin-bottom: 15px;
        position: relative;
    }
    
    .contract-content {
        line-height: 1.6;
        color: #495057;
    }
    
    /* Signature line animation */
    .signature-line {
        display: block;
        width: 100%;
        height: 1px;
        background-color: #000;
        margin: 15px 0;
        position: relative;
    }
    
    .signature-line::after {
        content: "";
        position: absolute;
        top: 0;
        left: 0;
        width: 0;
        height: 100%;
        background-color: #1a3c5a;
        transition: width 1.5s ease;
    }
    
    .signature-line:hover::after {
        width: 100%;
    }
    
    /* Party labels in contracts */
    .party-label {
        display: inline-block;
        background-color: rgba(26, 60, 90, 0.1);
        padding: 3px 8px;
        border-radius: 4px;
        font-weight: 500;
        color: #1a3c5a;
        margin-right: 5px;
        transition: all 0.2s;
    }
    
    .party-label:hover {
        background-color: rgba(26, 60, 90, 0.2);
        transform: translateY(-2px);
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
    }
    
    /* Enhanced form inputs for legal information */
    .legal-form-group {
        margin-bottom: 20px;
        position: relative;
    }
    
    .legal-form-label {
        display: block;
        margin-bottom: 8px;
        font-weight: 500;
        color: #1a3c5a;
        transition: all 0.3s;
    }
    
    .legal-form-input {
        width: 100%;
        padding: 12px 15px;
        border: 2px solid #e1e5eb;
        border-radius: 8px;
        font-size: 16px;
        transition: all 0.3s;
    }
    
    .legal-form-input:focus {
        border-color: #1a3c5a;
        box-shadow: 0 0 0 2px rgba(26, 60, 90, 0.2);
        outline: none;
    }
    
    .legal-form-input:focus + .legal-form-label {
        color: #2a5a85;
        transform: translateY(-25px) scale(0.8);
    }
    
    /* Loading states and transitions */
    .analysis-loading {
        display: flex;
        align-items: center;
        justify-content: center;
        padding: 20px;
    }
    
    .loading-dot {
        width: 10px;
        height: 10px;
        background-color: #1a3c5a;
        border-radius: 50%;
        margin: 0 5px;
        animation: bounce 1.5s infinite ease-in-out;
    }
    
    .loading-dot:nth-child(1) {animation-delay: 0s;}
    .loading-dot:nth-child(2) {animation-delay: 0.2s;}
    .loading-dot:nth-child(3) {animation-delay: 0.4s;}
    
    @keyframes bounce {
        0%, 80%, 100% {transform: scale(0);}
        40% {transform: scale(1);}
    }
    
    /* Help tooltips for legal terms */
    .legal-tooltip {
        position: relative;
        display: inline-block;
        border-bottom: 1px dotted #1a3c5a;
        cursor: help;
    }
    
    .legal-tooltip .tooltip-text {
        visibility: hidden;
        width: 300px;
        background-color: #1a3c5a;
        color: #fff;
        text-align: center;
        border-radius: 6px;
        padding: 10px;
        position: absolute;
        z-index: 1;
        bottom: 125%;
        left: 50%;
        transform: translateX(-50%);
        opacity: 0;
        transition: opacity 0.3s;
        box-shadow: 0 5px 15px rgba(0,0,0,0.2);
    }
    
    .legal-tooltip:hover .tooltip-text {
        visibility: visible;
        opacity: 1;
    }
    
    /* Evidence and Exhibits styling */
    .exhibit-container {
        border-left: 3px solid #1a3c5a;
        padding-left: 15px;
        margin: 20px 0;
        background-color: rgba(26, 60, 90, 0.05);
        padding: 15px;
        border-radius: 0 8px 8px 0;
        transition: all 0.3s;
    }
    
    .exhibit-container:hover {
        background-color: rgba(26, 60, 90, 0.1);
        transform: translateX(5px);
    }
    
    .exhibit-label {
        font-weight: 700;
        color: #1a3c5a;
        display: block;
        margin-bottom: 10px;
    }
    
    .exhibit-content {
        font-style: italic;
        color: #495057;
    }
    
    /* Contract template with elegant styling - fixed layout */
    .contract-template {
        padding: 30px;
        background-color: #f8f8f8;
        background-image: linear-gradient(45deg, #f9f9f9 25%, transparent 25%, transparent 75%, #f9f9f9 75%, #f9f9f9), 
                        linear-gradient(45deg, #f9f9f9 25%, transparent 25%, transparent 75%, #f9f9f9 75%, #f9f9f9);
        background-size: 20px 20px;
        background-position: 0 0, 10px 10px;
        border: 1px solid #ddd;
        border-radius: 8px;
        font-family: 'Playfair Display', serif;
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
        margin-bottom: 25px;
        width: 100%;
        box-sizing: border-box;
        overflow: auto;
    }

    /* Fix contract section styling */
    .contract-section {
        margin-bottom: 20px;
        padding: 15px;
        border-radius: 8px;
        background-color: rgba(255, 255, 255, 0.9);
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        position: relative;
        overflow: hidden;
        transition: all 0.3s;
        width: 100%;
        box-sizing: border-box;
    }

    /* Enhanced loading indicator */
    .loading-container {
        width: 100%;
        padding: 30px;
        text-align: center;
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
    }

    .loading-dot {
        width: 20px;
        height: 20px;
        background: linear-gradient(135deg, #1a3c5a 0%, #2a5a85 100%);
        border-radius: 50%;
        margin: 0 10px;
        display: inline-block;
        animation: bounce 1.5s infinite ease-in-out;
    }

    @keyframes bounce {
        0%, 80%, 100% {transform: translateY(0) scale(0.8);}
        40% {transform: translateY(-20px) scale(1);}
    }
    </style>
    """, unsafe_allow_html=True)

INDIAN_LANGUAGES = {
    "English": "en",
    "Hindi (हिन्दी)": "hi",
    "Tamil (தமிழ்)": "ta",
    "Telugu (తెలుగు)": "te",
    "Kannada (ಕನ್ನಡ)": "kn",
    "Malayalam (മലയാളം)": "ml",
    "Marathi (मराठी)": "mr",
    "Gujarati (ગુજરાતી)": "gu",
    "Bengali (বাংলা)": "bn",
    "Punjabi (ਪੰਜਾਬੀ)": "pa",
    "Odia (ଓଡ଼ିଆ)": "or",
    "Urdu (اردو)": "ur"
}

def translate_text(text, target_lang):
    """Translate text into the selected Indian language with error handling."""
    if target_lang == "en" or not text:
        return text
    
    try:
        # Split long text into smaller chunks to avoid translation errors
        max_chunk_length = 5000
        if len(text) > max_chunk_length:
            chunks = []
            for i in range(0, len(text), max_chunk_length):
                chunks.append(text[i:i+max_chunk_length])
            
            translated_chunks = []
            for chunk in chunks:
                translated_chunks.append(translator.translate(chunk, dest=target_lang).text)
            
            return "".join(translated_chunks)
        else:
            return translator.translate(text, dest=target_lang).text
    except Exception as e:
        print(f"Translation error: {str(e)}")
        return text  # Return original text if translation fails



def translate_ui_text(text_dict, target_lang):
    """Translate a dictionary of UI text elements to target language."""
    if target_lang == "en":
        return text_dict
    
    translated_dict = {}
    for key, text in text_dict.items():
        translated_dict[key] = translate_text(text, target_lang)
    
    return translated_dict

# Dictionary of UI text for translation
UI_TEXT = {
    "app_title": "JUSTISENSE",
    "chat_tab": "Chat",
    "document_analysis_tab": "Document Analysis",
    "contract_generator_tab": "Contract Generator",
    "history_tab": "History",
    "settings": "Settings & Tools",
    "language": "Choose Language:",
    "document_processing": "Document Processing",
    "upload_docs": "Upload your PDF Files",
    "process_button": "Process Documents",
    "processing_complete": "Processing complete. Ready for Q&A.",
    "ask_question": "Ask about your legal documents",
    "question_placeholder": "Ask your question...",
    "document_preview": "Document Preview:",
    "analyze_button": "Analyze Document",
    "key_info": "Key Information",
    "timeline": "Timeline of Events",
    "risk_assessment": "Risk Assessment",
    "legal_entities": "Legal Entities",
    "obligations": "Obligations Summary",
    "timeline_analysis": "Timeline Analysis",
    "legal_entities_info": "Legal Entities and Key Information",
    "select_contract": "Select contract type:",
    "employment_agreement": "Employment Agreement",
    "lease_agreement": "Lease Agreement",
    "nda": "Non-Disclosure Agreement (NDA)",
    "custom_contract": "Custom Contract",
    "employer_name": "Employer Name:",
    "employer_address": "Employer Address:",
    "agreement_date": "Agreement Date:",
    "employee_name": "Employee Name:",
    "employee_address": "Employee Address:",
    "position_title": "Position Title:",
    "landlord_name": "Landlord Name:",
    "landlord_address": "Landlord Address:",
    "tenant_name": "Tenant Name:",
    "tenant_address": "Tenant Address:",
    "property_description": "Property Description:",
    "party1_name": "First Party Name:",
    "party1_address": "First Party Address:",
    "party2_name": "Second Party Name:",
    "party2_address": "Second Party Address:",
    "purpose_disclosure": "Purpose of Disclosure:",
    "contract_description": "Describe the contract you need:",
    "parties_involved": "List parties involved (separated by comma):",
    "governing_law": "Governing Law:",
    "generate_contract": "Generate Contract",
    "generated_contract": "Generated Contract",
    "preview": "Preview:",
    "download_txt": "Download Text",
    "download_docx": "Download Word",
    "download_pdf": "Download PDF",
    "chat_history": "Chat History",
    "clear_history": "Clear History",
    "history_cleared": "Chat history cleared",
    "error_processing": "Error processing your question:",
    "try_again": "Try uploading and processing your documents again.",
    "upload_first": "Please upload and process documents first to enable analysis",
    "no_history": "No chat history found",
    "document_enhancement": "Document Enhancement",
    "enhancement_input": "Choose how to provide the document:",
    "use_processed_doc": "Use processed document",
    "upload_new_doc": "Upload new document",
    "paste_text": "Paste text",
    "using_processed_doc": "Using the processed document.",
    "no_processed_doc": "No processed document found. Please upload and process a document first.",
    "upload_enhance_doc": "Upload document to enhance",
    "document_loaded": "Document loaded successfully.",
    "document_load_error": "Error loading document:",
    "paste_document_text": "Paste the document text here:",
    "preview_original": "Preview Original Document",
    "enhancement_options": "Enhancement Options",
    "legal_formatting": "Apply legal formatting",
    "clause_numbering": "Add clause numbering",
    "term_definition": "Define legal terms",
    "citation_format": "Format citations",
    "clarity_improvement": "Improve clarity",
    "risk_identification": "Identify risks",
    "error_correction": "Correct errors",
    "consistency_check": "Check for consistency",
    "selected_enhancements": "Selected Enhancements",
    "no_enhancements_selected": "No enhancements selected.",
    "document_type": "Document Type",
    "contract": "Contract",
    "legal_brief": "Legal Brief",
    "court_filing": "Court Filing",
    "memo": "Memo",
    "corporate_document": "Corporate Document",
    "other": "Other",
    "additional_instructions": "Additional Instructions",
    "enhance_document": "Enhance Document",
    "enhancing_document": "Enhancing document...",
    "enhanced_document": "Enhanced Document",
    "enhancements_applied": "Enhancements Applied",
    "preview_enhanced": "Preview Enhanced Document",
    "docx_download_error": "Error downloading DOCX. Please try again.",
    "pdf_download_error": "Error downloading PDF. Please try again.",
    "show_comparison": "Show Comparison",
    "original_document": "Original Document",
    "enhancement_error": "Error enhancing document:",
    "provide_document": "Please provide a document to enhance."
}

# Add these entries to your UI_TEXT dictionary:

# Update the UI_TEXT dictionary with new entries
UI_TEXT.update({
    # Document Enhancement tab
    "document_enhancement": "Document Enhancement",
    "enhancement_input": "Select document input method:",
    "use_processed_doc": "Use already processed document",
    "upload_new_doc": "Upload a new document",
    "paste_text": "Paste document text",
    "no_processed_doc": "No document has been processed yet. Please upload a document first.",
    "using_processed_doc": "Using the previously processed document",
    "upload_enhance_doc": "Upload a document to enhance",
    "paste_document_text": "Paste your document text here",
    "document_loaded": "Document loaded successfully",
    "document_load_error": "Error loading document:",
    "docx_import_error": "Python-docx module not found. Please install it to process DOCX files.",
    "preview_original": "Preview Original Document",
    "enhancement_options": "Enhancement Options",
    "legal_formatting": "Legal Formatting",
    "clause_numbering": "Clause Numbering and Structure",
    "term_definition": "Define Legal Terms",
    "citation_format": "Format Citations",
    "clarity_improvement": "Improve Clarity and Readability",
    "risk_identification": "Identify Legal Risks",
    "error_correction": "Correct Grammar and Language Errors",
    "consistency_check": "Check Term Consistency",
    "selected_enhancements": "Selected Enhancements:",
    "no_enhancements_selected": "No enhancements selected. Please select at least one enhancement option.",
    "document_type": "Document Type:",
    "contract": "Contract",
    "legal_brief": "Legal Brief",
    "court_filing": "Court Filing",
    "memo": "Legal Memorandum",
    "corporate_document": "Corporate Document",
    "other": "Other",
    "additional_instructions": "Additional Instructions (Optional):",
    "enhance_document": "Enhance Document",
    "enhancing_document": "Enhancing document...",
    "enhanced_document": "Enhanced Document",
    "enhancements_applied": "Enhancements Applied:",
    "preview_enhanced": "Preview Enhanced Document",
    "show_comparison": "Show Side-by-Side Comparison",
    "original_document": "Original Document",
    "enhancement_error": "Error enhancing document:",
    "try_again": "Please try again. If the problem persists, try with a different document.",
    "provide_document": "Please provide a document to enhance",
    "docx_download_error": "Install python-docx for Word downloads",
    "pdf_download_error": "Install reportlab for PDF downloads"
})

UI_TEXT.update({
    # Add these entries for the Voice Chat tab
    "voice_chat_tab": "Voice Chat",
    "voice_language": "Select Language:",
    "start_recording": "Start Recording",
    "recording": "Recording... Speak now!",
    "processing": "Processing...",
    "your_input": "Your Input Text:",
    "gemini_response": "Legal Assistant's Response:",
    "download_audio": "Download Audio",
    "recording_duration": "Recording Duration (seconds):",
    "no_speech_detected": "No speech was detected. Please try again.",
    "api_key_missing": "API key is missing. Please check your environment variables.",
    "voice_chat_intro": "Speak in your preferred language to get legal assistance",
    "voice_processing_error": "Error processing voice: "
})

UI_TEXT.update({
    # Add these entries for the Legal Case Finder tab
    "case_finder_tab": "Case Finder",
    "case_finder_title": "Legal Case & News Finder",
    "case_finder_description": "Search for relevant legal cases and news articles by keyword",
    "search_keyword": "Enter a keyword to find related legal cases & news:",
    "search_button": "Search",
    "searching_message": "Searching for articles and cases related to '{}'...",
    "related_results": "Related Legal Cases & News",
    "no_results": "No relevant pages found. Try another keyword.",
    "enter_keyword": "Please enter a keyword.",
    "source": "Source:",
    "open_link": "Open Link",
    "packages_missing": "Missing required packages: {}. Please install selenium and webdriver-manager."
})

def get_pdf_text(pdf_docs):
    """Extract text from uploaded PDFs with unicode error handling."""
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    # Clean text of problematic unicode characters
                    page_text = clean_text(page_text)
                    text += page_text + "\n"
            except Exception as e:
                st.warning(f"Error extracting text from a page: {str(e)}")
    return text

def clean_text(text):
    """Clean text of problematic unicode characters."""
    if not text:
        return ""
    
    # Replace problematic unicode characters
    cleaned = ""
    for char in text:
        try:
            # Test if character can be encoded
            char.encode('utf-8')
            cleaned += char
        except UnicodeEncodeError:
            # Replace problematic character with a space or similar
            cleaned += " "
    return cleaned

def get_text_chunks(text):
    """Split text into chunks for vector storage."""
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    return text_splitter.split_text(text)

def get_vector_store(text_chunks):
    """Create FAISS vector store from text chunks with error handling."""
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    
    # Process in smaller batches to avoid issues with very large documents
    vector_store = None
    batch_size = 100  # Adjust based on your needs
    
    try:
        for i in range(0, len(text_chunks), batch_size):
            batch = text_chunks[i:i+batch_size]
            
            # Process this batch
            try:
                if vector_store is None:
                    vector_store = FAISS.from_texts(batch, embedding=embeddings)
                else:
                    batch_vectors = FAISS.from_texts(batch, embedding=embeddings)
                    vector_store.merge_from(batch_vectors)
                
                # Show progress
                progress = min(100, int((i + batch_size) / len(text_chunks) * 100))
                st.progress(progress / 100)
                
            except UnicodeEncodeError as e:
                st.warning(f"Unicode error in batch {i//batch_size + 1}. Cleaning batch and retrying.")
                # Clean the batch more aggressively
                cleaned_batch = [clean_text_aggressive(chunk) for chunk in batch]
                
                if vector_store is None:
                    vector_store = FAISS.from_texts(cleaned_batch, embedding=embeddings)
                else:
                    batch_vectors = FAISS.from_texts(cleaned_batch, embedding=embeddings)
                    vector_store.merge_from(batch_vectors)
        
        # Save the vector store
        if vector_store:
            vector_store.save_local("faiss_index")
            return True
        return False
    
    except Exception as e:
        st.error(f"Error creating vector store: {str(e)}")
        return False

def clean_text_aggressive(text):
    """More aggressive text cleaning for problematic chunks."""
    if not text:
        return ""
    
    # Replace all non-ASCII characters with spaces
    return ''.join(c if ord(c) < 128 else ' ' for c in text)

def get_conversational_chain():
    """Initialize Gemini conversational chain."""
    prompt_template = """
    Answer the question as detailed as possible from the provided context. 
    If the answer is not in the provided context, respond with "answer is not available in the context".
    
    Context:\n {context}\n
    Question: \n{question}\n
    Answer:
    """
    model = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    return load_qa_chain(model, chain_type="stuff", prompt=prompt)

# Update user_input function to use new styling
def user_input(user_question, target_lang):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

    # Ensure FAISS index exists
    if not os.path.exists("faiss_index"):
        st.error("No FAISS index found. Please upload and process PDFs first.")
        return

    # Load FAISS index safely
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    
    docs = new_db.similarity_search(user_question)
    chain = get_conversational_chain()
    
    with st.status("Processing your question...", expanded=True) as status:
        st.write("Searching through documents...")
        response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
        reply = response.get("output_text", "No response generated.")
        
        st.write("Translating response...")
        translated_reply = translate_text(reply, target_lang)
        status.update(label="Complete!", state="complete")
    
    # Display in custom message container
    st.markdown(
        f'<div class="chat-message user"><div class="message-content"><b>{translate_text("You", target_lang)}:</b> {user_question}</div></div>', 
        unsafe_allow_html=True
    )
    
    # Add a small delay for visual effect
    import time
    time.sleep(0.5)
    
    st.markdown(
        f'<div class="chat-message bot"><div class="message-content"><b>{translate_text("Legal Assistant", target_lang)}:</b> {translated_reply}</div></div>', 
        unsafe_allow_html=True
    )
    
    # Save to history
    save_chat_history(user_question, translated_reply)
    
    return reply

# Chat history management function
def save_chat_history(question, answer):
    """Save chat interaction to history file"""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    history_item = {
        "timestamp": timestamp,
        "question": question,
        "answer": answer
    }
    
    # Create directory if it doesn't exist
    os.makedirs("history", exist_ok=True)
    
    history_file = "history/chat_history.json"
    try:
        if os.path.exists(history_file):
            with open(history_file, "r") as f:
                history = json.load(f)
        else:
            history = []
            
        history.append(history_item)
        
        with open(history_file, "w") as f:
            json.dump(history, f, indent=4)
    except Exception as e:
        st.error(f"Error saving history: {str(e)}")

def get_chat_history():
    """Retrieve chat history from file"""
    history_file = "history/chat_history.json"
    if os.path.exists(history_file):
        try:
            with open(history_file, "r") as f:
                return json.load(f)
        except:
            return []
    return []

def display_chat_history():
    """Display chat history in the UI"""
    history = get_chat_history()
    
    if not history:
        st.info("No chat history found")
        return
        
    st.markdown("<div class='history-container'>", unsafe_allow_html=True)
    
    for item in reversed(history):  # Most recent first
        st.markdown(
            f"<div class='history-item'>"
            f"<small>{item['timestamp']}</small><br>"
            f"<strong>Q: {item['question']}</strong><br>"
            f"A: {item['answer'][:100]}{'...' if len(item['answer']) > 100 else ''}"
            f"</div>", 
            unsafe_allow_html=True
        )
    
    st.markdown("</div>", unsafe_allow_html=True)

def extract_timeline(document_text):
    """Extract a chronological timeline of events from legal document with proper formatting."""
    prompt = """
    Extract a chronological timeline of key events from this legal document.
    For each event, provide:
    1. Date (if available)
    2. Event description
    3. Relevant parties involved
    4. Page reference (if available)
    
    Format EXACTLY as a JSON array of objects with the following keys:
    [
      {
        "date": "YYYY-MM-DD or textual date",
        "description": "Description of the event",
        "parties": "Parties involved",
        "page": "Page reference"
      },
      ...
    ]
    
    Ensure your output is valid JSON only, with no additional text before or after.
    
    Document text:
    {text}
    """
    
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(prompt.format(text=document_text[:50000]))
        
        # Try to parse as JSON
        try:
            # Strip any extra text or formatting that might be around the JSON
            response_text = response.text.strip()
            
            # Find the start and end of what looks like a JSON array
            if '[' in response_text and ']' in response_text:
                start = response_text.find('[')
                end = response_text.rfind(']') + 1
                json_str = response_text[start:end]
                
                # Parse the JSON
                timeline_data = json.loads(json_str)
                return timeline_data
            else:
                return response.text
        except json.JSONDecodeError:
            # If not valid JSON, return the raw text
            return response.text
    except Exception as e:
        return f"Error extracting timeline: {str(e)}"

def extract_legal_entities(document_text):
    """Extract key legal entities and information from the document."""
    prompt = """
    Extract key legal entities and information from this document. Include:
    
    1. All named parties (individuals and organizations)
    2. Key dates and deadlines
    3. Monetary values and financial terms
    4. Property or asset descriptions
    5. Legal obligations and requirements
    
    Format the response as structured JSON with these categories as keys, and lists of extracted items as values.
    Only return the JSON, no additional text.
    
    Document text:
    {text}
    """
    
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(prompt.format(text=document_text[:50000]))
        
        # Parse the JSON response
        try:
            # Extract only the JSON part if there's other text
            response_text = response.text
            if '{' in response_text and '}' in response_text:
                start = response_text.find('{')
                end = response_text.rfind('}') + 1
                json_str = response_text[start:end]
                return json.loads(json_str)
            else:
                return {"error": "Could not extract structured data"}
        except json.JSONDecodeError:
            return {"error": "Invalid JSON response", "raw": response.text[:500]}
    except Exception as e:
        return {"error": f"Error extracting entities: {str(e)}"}

def display_timeline(timeline_data):
    """Display timeline in a visually appealing format."""
    if isinstance(timeline_data, list):
        # It's properly formatted JSON
        st.markdown("<div class='timeline-container'>", unsafe_allow_html=True)
        
        for i, event in enumerate(timeline_data):
            date = event.get('date', 'No date')
            description = event.get('description', 'No description')
            parties = event.get('parties', 'Not specified')
            page = event.get('page', 'N/A')
            
            position = "left" if i%2==0 else "right"
            
            # Fix timeline CSS layout issues
            st.markdown(
                f"""
                <div class='timeline-item {position}' style='clear: both;'>
                    <div class='timeline-content' style='width: 100%; box-sizing: border-box;'>
                        <div class='timeline-date'>{date}</div>
                        <h4>{description}</h4>
                        <p><strong>Parties:</strong> {parties}</p>
                        <p><strong>Page:</strong> {page}</p>
                    </div>
                </div>
                """,
                unsafe_allow_html=True
            )
        
        st.markdown("</div>", unsafe_allow_html=True)
    else:
        # It's raw text
        st.markdown("### Timeline")
        st.text_area("Timeline extraction result:", timeline_data, height=300)

def generate_contract_template(context_info):
    """Generate a contract template based on user input."""
    contract_types = {
        "employment": """
        EMPLOYMENT AGREEMENT

        This Employment Agreement (the "Agreement") is made and entered into as of [DATE], by and between:

        EMPLOYER: [EMPLOYER_NAME], with its principal place of business at [EMPLOYER_ADDRESS] ("Employer")

        EMPLOYEE: [EMPLOYEE_NAME], residing at [EMPLOYEE_ADDRESS] ("Employee")

        1. POSITION AND DUTIES
        [POSITION_DETAILS]

        2. TERM OF EMPLOYMENT
        [TERM_DETAILS]

        3. COMPENSATION AND BENEFITS
        [COMPENSATION_DETAILS]

        4. TERMINATION
        [TERMINATION_DETAILS]

        5. CONFIDENTIALITY
        [CONFIDENTIALITY_DETAILS]

        6. NON-COMPETE
        [NON_COMPETE_DETAILS]

        7. GOVERNING LAW
        This Agreement shall be governed by and construed in accordance with the laws of [JURISDICTION].

        IN WITNESS WHEREOF, the parties hereto have executed this Agreement as of the date first above written.

        ________________________
        [EMPLOYER_NAME]
        By: [AUTHORIZED_SIGNATORY]
        Title: [TITLE]

        ________________________
        [EMPLOYEE_NAME]
        """,
        
        "lease": """
        LEASE AGREEMENT

        This Lease Agreement (the "Lease") is made and entered into as of [DATE], by and between:

        LANDLORD: [LANDLORD_NAME], with an address of [LANDLORD_ADDRESS] ("Landlord")

        TENANT: [TENANT_NAME], with an address of [TENANT_ADDRESS] ("Tenant")

        1. PREMISES
        [PROPERTY_DETAILS]

        2. TERM
        [LEASE_TERM]

        3. RENT
        [RENT_DETAILS]

        4. SECURITY DEPOSIT
        [DEPOSIT_DETAILS]

        5. UTILITIES AND SERVICES
        [UTILITIES_DETAILS]

        6. MAINTENANCE AND REPAIRS
        [MAINTENANCE_DETAILS]

        7. DEFAULT
        [DEFAULT_DETAILS]

        8. GOVERNING LAW
        This Lease shall be governed by and construed in accordance with the laws of [JURISDICTION].

        IN WITNESS WHEREOF, the parties hereto have executed this Lease as of the date first above written.

        ________________________
        [LANDLORD_NAME]

        ________________________
        [TENANT_NAME]
        """,
        
        "nda": """
        NON-DISCLOSURE AGREEMENT

        This Non-Disclosure Agreement (the "Agreement") is made and entered into as of [DATE], by and between:

        PARTY 1: [PARTY1_NAME], with its principal place of business at [PARTY1_ADDRESS] ("Disclosing Party")

        PARTY 2: [PARTY2_NAME], with its principal place of business at [PARTY2_ADDRESS] ("Receiving Party")

        1. PURPOSE
        [PURPOSE_DETAILS]

        2. DEFINITION OF CONFIDENTIAL INFORMATION
        [CONFIDENTIAL_INFO_DETAILS]

        3. OBLIGATIONS OF RECEIVING PARTY
        [OBLIGATIONS_DETAILS]

        4. TERM
        [TERM_DETAILS]

        5. REMEDIES
        [REMEDIES_DETAILS]

        6. GOVERNING LAW
        This Agreement shall be governed by and construed in accordance with the laws of [JURISDICTION].

        IN WITNESS WHEREOF, the parties hereto have executed this Agreement as of the date first above written.

        ________________________
        [PARTY1_NAME]
        By: [AUTHORIZED_SIGNATORY1]
        Title: [TITLE1]

        ________________________
        [PARTY2_NAME]
        By: [AUTHORIZED_SIGNATORY2]
        Title: [TITLE2]
        """
    }
    
    # Fill template with context
    contract_type = context_info.get("contract_type", "employment")
    template = contract_types.get(contract_type, contract_types["employment"])
    
    # Basic replacements
    for key, value in context_info.items():
        placeholder = f"[{key.upper()}]"
        if placeholder in template:
            template = template.replace(placeholder, value)
    
    # Generate remaining sections using AI
    remaining_placeholders = [p for p in re.findall(r'\[(.*?)\]', template)]
    if remaining_placeholders:
        try:
            # Generate content for remaining placeholders
            model = genai.GenerativeModel('gemini-1.5-flash')
            for placeholder in remaining_placeholders:
                prompt = f"Generate appropriate legal language for the {placeholder} section of a {contract_type} contract."
                response = model.generate_content(prompt)
                template = template.replace(f"[{placeholder}]", response.text)
        except Exception as e:
            pass  # Keep placeholder if generation fails
    
    return template

def generate_structured_contract(template_type, context_info):
    """Generate a structured contract based on template and context information."""
    # Basic contract templates with formatting
    templates = {
        "employment": {
            "title": "EMPLOYMENT AGREEMENT",
            "sections": [
                {"heading": "PARTIES", "content": """
This Employment Agreement (the "Agreement") is made and entered into as of {date}, by and between:

EMPLOYER: {employer_name}, with its principal place of business at {employer_address} ("Employer")

EMPLOYEE: {employee_name}, residing at {employee_address} ("Employee")
                """},
                {"heading": "POSITION AND DUTIES", "content": """
Employer hereby employs Employee as {position}, and Employee accepts such employment, on the terms and conditions set forth in this Agreement. Employee shall perform such duties as are customarily performed by others in similar positions at Employer, as well as such other duties as may be assigned from time to time by Employer.
                """},
                {"heading": "TERM OF EMPLOYMENT", "content": """
The term of this Agreement shall commence on {date} and shall continue until terminated in accordance with the provisions of this Agreement.
                """},
                {"heading": "COMPENSATION AND BENEFITS", "content": """
Employee shall be paid a base salary of [SALARY_AMOUNT] per [PAY_PERIOD], subject to applicable withholdings and deductions.

Employee shall be eligible to participate in Employer's employee benefit plans as may be in effect from time to time.
                """},
                {"heading": "GOVERNING LAW", "content": """
This Agreement shall be governed by and construed in accordance with the laws of {jurisdiction}.
                """},
                {"heading": "SIGNATURES", "content": """
IN WITNESS WHEREOF, the parties hereto have executed this Agreement as of the date first above written.

________________________
{employer_name}

________________________
{employee_name}
                """}
            ]
        },
        "nda": {
            "title": "NON-DISCLOSURE AGREEMENT",
            "sections": [
                {"heading": "PARTIES", "content": """
This Non-Disclosure Agreement (the "Agreement") is made and entered into as of {date}, by and between:

PARTY 1: {party1_name}, with its principal place of business at {party1_address} ("Disclosing Party")

PARTY 2: {party2_name}, with its principal place of business at {party2_address} ("Receiving Party")
                """},
                {"heading": "PURPOSE", "content": """
{purpose_details}
                """},
                {"heading": "DEFINITION OF CONFIDENTIAL INFORMATION", "content": """
"Confidential Information" means any information disclosed by Disclosing Party to Receiving Party, either directly or indirectly, in writing, orally or by inspection of tangible objects, which is designated as "Confidential," "Proprietary" or some similar designation, or that should reasonably be understood to be confidential given the nature of the information and the circumstances of disclosure.
                """},
                {"heading": "OBLIGATIONS OF RECEIVING PARTY", "content": """
Receiving Party shall hold the Confidential Information in strict confidence and shall not disclose such Confidential Information to any third party. Receiving Party shall protect the confidentiality of the Confidential Information using at least the same degree of care it uses to protect its own confidential information, but no less than reasonable care.
                """},
                {"heading": "GOVERNING LAW", "content": """
This Agreement shall be governed by and construed in accordance with the laws of {jurisdiction}.
                """},
                {"heading": "SIGNATURES", "content": """
IN WITNESS WHEREOF, the parties hereto have executed this Agreement as of the date first above written.

________________________
{party1_name}

________________________
{party2_name}
                """}
            ]
        },
        "lease": {
            "title": "LEASE AGREEMENT",
            "sections": [
                {"heading": "PARTIES", "content": """
This Lease Agreement (the "Lease") is made and entered into as of {date}, by and between:

LANDLORD: {landlord_name}, with an address of {landlord_address} ("Landlord")

TENANT: {tenant_name}, with an address of {tenant_address} ("Tenant")
                """},
                {"heading": "PREMISES", "content": """
{property_details}
                """},
                {"heading": "TERM", "content": """
The term of this Lease shall commence on {date} and shall continue for a period of [LEASE_DURATION], unless sooner terminated in accordance with the provisions of this Lease.
                """},
                {"heading": "RENT", "content": """
Tenant shall pay to Landlord as rent for the Premises the sum of [RENT_AMOUNT] per month, payable in advance on the first day of each month during the term of this Lease.
                """},
                {"heading": "GOVERNING LAW", "content": """
This Lease shall be governed by and construed in accordance with the laws of {jurisdiction}.
                """},
                {"heading": "SIGNATURES", "content": """
IN WITNESS WHEREOF, the parties hereto have executed this Lease as of the date first above written.

________________________
{landlord_name}

________________________
{tenant_name}
                """}
            ]
        }
    }
    
    # Select the appropriate template
    template = templates.get(template_type)
    if not template:
        return "Template type not found."
    
    # Build the contract
    contract = template["title"] + "\n\n"
    
    for section in template["sections"]:
        contract += section["heading"] + "\n\n"
        # Format the content with context info
        content = section["content"]
        for key, value in context_info.items():
            if isinstance(value, str) and value.strip():
                content = content.replace("{" + key + "}", value)
        contract += content.strip() + "\n\n"
    
    # Fill in any remaining placeholders with AI-generated content
    placeholders = re.findall(r'\[(.*?)\]', contract)
    if placeholders:
        try:
            model = genai.GenerativeModel('gemini-1.5-flash')
            for placeholder in placeholders:
                # Add instructions to avoid markdown formatting
                prompt = f"""Generate appropriate legal language for {placeholder} in a {template_type} contract.
                
                IMPORTANT: DO NOT include any markdown formatting symbols like asterisks (*), 
                greater-than symbols (>), or bullet points. Return ONLY clean, plain text suitable 
                for direct inclusion in a legal document. Do not include 'Option 1/2/3' language or disclaimers.
                Keep your response concise and formal in legal language (2-3 sentences maximum)."""
                
                response = model.generate_content(prompt)
                
                # Clean up the response to remove any remaining markdown
                cleaned_response = response.text
                cleaned_response = re.sub(r'[*>•\-]', '', cleaned_response)
                cleaned_response = re.sub(r'Option \d+[:\)]', '', cleaned_response)
                cleaned_response = cleaned_response.replace('', '').strip()
                
                contract = contract.replace(f"[{placeholder}]", cleaned_response)
        except Exception as e:
            # Keep placeholders if generation fails
            print(f"Error generating content: {str(e)}")
    
    return contract

def download_document(document_text, filename="legal_document.txt"):
    """Create a download link for a document with proper formatting"""
    file_ext = filename.split('.')[-1].lower()
    
    if file_ext == 'docx':
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Create a professional Word document
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add title
            title = doc.add_heading('', 0)
            title_run = title.add_run("LEGAL DOCUMENT")
            title_run.bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add content with proper formatting
            paragraphs = document_text.split('\n')
            for para in paragraphs:
                if para.strip():
                    if para.strip().isupper() and len(para.strip()) < 50:  # Likely a heading
                        p = doc.add_heading(para.strip(), level=1)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p = doc.add_paragraph(para)
                        # Check if this is a signature line
                        if "" in para:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Save to BytesIO object
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            file_bytes = docx_bytes.getvalue()
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        except ImportError:
            # Fall back to text if python-docx is not installed
            file_bytes = document_text.encode()
            mime_type = "text/plain"
            
    elif file_ext == 'pdf':
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch
            
            # Create a PDF with professional formatting
            buffer = io.BytesIO()
            
            # Set up the document
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Define styles - FIX: Don't redefine existing styles
            styles = getSampleStyleSheet()
            
            # Create a custom heading style instead of modifying the built-in one
            heading_style = ParagraphStyle(
                name='CustomHeading',
                parent=styles['Heading1'],
                alignment=1,  # Center alignment
                spaceAfter=12
            )
            
            body_style = ParagraphStyle(
                name='CustomBody',
                parent=styles['BodyText'],
                spaceBefore=6,
                spaceAfter=6
            )
            
            # Process content
            content = []
            content.append(Paragraph("LEGAL DOCUMENT", heading_style))
            content.append(Spacer(1, 0.25*inch))
            
            # Break text into paragraphs
            paragraphs = document_text.split('\n')
            for para in paragraphs:
                if not para.strip():
                    content.append(Spacer(1, 0.1*inch))
                    continue
                    
                if para.strip().isupper() and len(para.strip()) < 50:  # Likely a heading
                    content.append(Spacer(1, 0.1*inch))
                    content.append(Paragraph(para, heading_style))
                else:
                    # Escape any special characters in the paragraph text
                    para_text = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    content.append(Paragraph(para_text, body_style))
            
            # Build the PDF
            doc.build(content)
            
            buffer.seek(0)
            file_bytes = buffer.getvalue()
            mime_type = "application/pdf"
        except Exception as e:
            # Fall back to text with error information
            file_bytes = f"PDF generation failed: {str(e)}\n\n{document_text}".encode()
            mime_type = "text/plain"
    else:
        # Default to text file
        file_bytes = document_text.encode()
        mime_type = "text/plain"
    
    b64 = base64.b64encode(file_bytes).decode()
    href = f'<a href="data:{mime_type};base64,{b64}" download="{filename}" class="download-button">Download {filename}</a>'
    return href

def show_contract_preview(context_info, contract_type):
    """Show interactive contract preview as user inputs data"""
    try:
        # Check if we have the minimum required information
        has_required_info = False
        if contract_type == "employment" and context_info.get("employer_name"):
            has_required_info = True
        elif contract_type == "nda" and context_info.get("party1_name"):
            has_required_info = True
        elif contract_type == "lease" and context_info.get("landlord_name"):
            has_required_info = True
            
        if not has_required_info:
            return
        
        st.markdown("### Live Contract Preview")
        
        # Show only the first part of the contract in the preview
        if contract_type == "employment":
            preview_html = f"""
            <div class='contract-section'>
                <h4 class='contract-heading'>EMPLOYMENT AGREEMENT</h4>
                <div class='contract-content'>
                    This Employment Agreement (the "Agreement") is made and entered into as of {context_info.get("date", "[DATE]")}, by and between:
                    
                    <p><span class='party-label'>EMPLOYER</span>: <strong>{context_info.get("employer_name", "[EMPLOYER NAME]")}</strong>, 
                    with its principal place of business at {context_info.get("employer_address", "[ADDRESS]")} ("Employer")</p>
                    
                    <p><span class='party-label'>EMPLOYEE</span>: <strong>{context_info.get("employee_name", "[EMPLOYEE NAME]")}</strong>, 
                    residing at {context_info.get("employee_address", "[ADDRESS]")} ("Employee")</p>
                    
                    <p>Employer hereby employs Employee as <strong>{context_info.get("position", "[POSITION]")}</strong>, and Employee accepts such employment, on the terms and conditions set forth in this Agreement.</p>
                </div>
            </div>
            """
        elif contract_type == "nda":
            preview_html = f"""
            <div class='contract-section'>
                <h4 class='contract-heading'>NON-DISCLOSURE AGREEMENT</h4>
                <div class='contract-content'>
                    This Non-Disclosure Agreement (the "Agreement") is made and entered into as of {context_info.get("date", "[DATE]")}, by and between:
                    
                    <p><span class='party-label'>PARTY 1</span>: <strong>{context_info.get("party1_name", "[PARTY 1 NAME]")}</strong>, 
                    with its principal place of business at {context_info.get("party1_address", "[ADDRESS]")} ("Disclosing Party")</p>
                    
                    <p><span class='party-label'>PARTY 2</span>: <strong>{context_info.get("party2_name", "[PARTY 2 NAME]")}</strong>, 
                    with its principal place of business at {context_info.get("party2_address", "[ADDRESS]")} ("Receiving Party")</p>
                    
                    <p><strong>PURPOSE:</strong> {context_info.get("purpose_details", "[PURPOSE OF DISCLOSURE]")}</p>
                </div>
            </div>
            """
        elif contract_type == "lease":
            preview_html = f"""
            <div class='contract-section'>
                <h4 class='contract-heading'>LEASE AGREEMENT</h4>
                <div class='contract-content'>
                    This Lease Agreement (the "Lease") is made and entered into as of {context_info.get("date", "[DATE]")}, by and between:
                    
                    <p><span class='party-label'>LANDLORD</span>: <strong>{context_info.get("landlord_name", "[LANDLORD NAME]")}</strong>, 
                    with an address of {context_info.get("landlord_address", "[ADDRESS]")} ("Landlord")</p>
                    
                    <p><span class='party-label'>TENANT</span>: <strong>{context_info.get("tenant_name", "[TENANT NAME]")}</strong>, 
                    with an address of {context_info.get("tenant_address", "[ADDRESS]")} ("Tenant")</p>
                    
                    <p><strong>PREMISES:</strong> {context_info.get("property_details", "[PROPERTY DETAILS]")}</p>
                </div>
            </div>
            """
        
        # Wrap the preview in the contract template container
        st.markdown(f"<div class='contract-template'>{preview_html}</div>", unsafe_allow_html=True)
        
    except Exception as e:
        st.error(f"Error showing preview: {str(e)}")

# Add a custom loading spinner component
def show_custom_spinner(message="Processing..."):
    with st.container():
        col1, col2, col3 = st.columns([1, 3, 1])
        with col2:
            st.markdown("""
            <div class="loading-container" style="text-align: center;">
                <div style="display: flex; justify-content: center; margin-bottom: 20px;">
                    <div class="loading-dot" style="animation-delay: 0s;"></div>
                    <div class="loading-dot" style="animation-delay: 0.2s;"></div>
                    <div class="loading-dot" style="animation-delay: 0.4s;"></div>
                </div>
                <p>{}</p>
            </div>
            """.format(message), unsafe_allow_html=True)

def record_audio(duration=10, sample_rate=16000):
    """Record audio from microphone for specified duration"""
    try:
        import pyaudio
        import numpy as np
        import soundfile as sf
        
        CHUNK = 1024
        FORMAT = pyaudio.paFloat32
        CHANNELS = 1

        p = pyaudio.PyAudio()
        stream = p.open(format=FORMAT, channels=CHANNELS, rate=sample_rate, input=True, frames_per_buffer=CHUNK)

        frames = []
        
        with st.status("Recording... Speak now!", expanded=True) as status:
            for i in range(0, int(sample_rate / CHUNK * duration)):
                data = stream.read(CHUNK)
                frames.append(np.frombuffer(data, dtype=np.float32))
                # Update progress
                progress = (i+1)/(int(sample_rate / CHUNK * duration))
                status.update(label=f"Recording... {int(progress*100)}%", state="running")
            
            status.update(label="Recording complete", state="complete")

        stream.stop_stream()
        stream.close()
        p.terminate()

        audio_array = np.concatenate(frames, axis=0)
        temp_file = "temp_recording.wav"
        sf.write(temp_file, audio_array, sample_rate)
        return temp_file
    
    except Exception as e:
        st.error(f"Error processing voice: {str(e)}")
        return None

def speech_to_text(audio_file, language_code):
    """Convert speech to text using Google Speech Recognition"""
    try:
        import speech_recognition as sr
        
        # Create recognizer instance
        r = sr.Recognizer()
        
        # Load the audio file
        with sr.AudioFile(audio_file) as source:
            audio_data = r.record(source)
            
            # Use Google Speech Recognition
            text = r.recognize_google(audio_data, language=language_code)
            return text
    
    except ImportError:
        st.error("Missing speech_recognition package. Install with: pip install SpeechRecognition")
        return None
    except Exception as e:
        st.error(f"Error processing voice: {str(e)}")
        return None

def text_to_speech(text, language_code, output_path="output_speech.mp3"):
    """Convert text to speech using ElevenLabs"""
    try:
        import requests
        
        # Use a default key if environment variable isn't set
        elevenlabs_api_key = os.getenv("ELEVENLABS_API_KEY", "sk_8153cf14f63223696a6e18d5d9ea9f64a6081f3f0ee68bcd")
        
        # Default voice ID (multilingual)
        voice_id = "21m00Tcm4TlvDq8ikWAM"
        
        headers = {"xi-api-key": elevenlabs_api_key, "Content-Type": "application/json"}
        data = {
            "text": text,
            "model_id": "eleven_multilingual_v2",  # Multilingual model for Indian languages
            "voice_settings": {"stability": 0.5, "similarity_boost": 0.5}
        }
        
        # Log request information for debugging (optional)
        print(f"Making request to ElevenLabs API with text length: {len(text)}")
        
        response = requests.post(
            f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}", 
            json=data, 
            headers=headers
        )
        
        if response.status_code == 200:
            with open(output_path, "wb") as f:
                f.write(response.content)
            return output_path
        else:
            error_message = f"TTS Error: Status {response.status_code}"
            try:
                error_details = response.json()
                error_message += f" - {error_details.get('detail', '')}"
            except:
                error_message += f" - {response.text[:  100]}"
            
            st.error(error_message)
            print(error_message)  # Log to console for debugging
            return None
    
    except Exception as e:
        error_message = f"Error processing voice: {str(e)}"
        st.error(error_message)
        print(error_message)  # Log to console for debugging
        return None

def fallback_text_to_speech(text, language_code, output_path="output_speech.mp3"):
    """Fallback TTS using gTTS when ElevenLabs fails"""
    try:
        from gtts import gTTS
        
        # Map language code if needed (adjust according to your needs)
        google_lang = language_code.split("-")[0] if "-" in language_code else language_code
        
        tts = gTTS(text=text, lang=google_lang, slow=False)
        tts.save(output_path)
        return output_path
        
    except Exception as e:
        st.error(f"Fallback TTS error: {str(e)}")
        return None

def generate_related_queries(keyword):
    """Generate broader search queries for legal research"""
    variations = [
        f"{keyword} court case news",
        f"{keyword} legal judgment pdf",
        f"{keyword} supreme court ruling",
        f"{keyword} lawsuit history",
        f"{keyword} government legal documents",
        f"{keyword} case law reference",
        f"{keyword} verdict and appeals",
        f"{keyword} high court decision",
        f"{keyword} latest legal updates",
        f"{keyword} legal case study",
        f"{keyword} legal precedents",
        f"{keyword} litigation process",
        f"{keyword} law enforcement cases",
        f"{keyword} judicial decisions",
        f"{keyword} supreme court review",
        f"{keyword} appeals court ruling",
    ]
    return random.sample(variations, min(len(variations), 7))  # Select 7 queries randomly

def scrape_links(keyword, max_results=15):
    """Scrape legal case links for multiple queries"""
    queries = generate_related_queries(keyword)
    all_links = set()

    options = Options()
    options.add_argument("--headless")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")

    try:
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)

        for query in queries:
            search_url = f"https://www.bing.com/search?q={urllib.parse.quote(query)}"
            driver.get(search_url)
            time.sleep(2)

            search_results = driver.find_elements(By.CSS_SELECTOR, "li.b_algo")
            for result in search_results:
                try:
                    link_element = result.find_element(By.TAG_NAME, "a")
                    link = link_element.get_attribute("href")
                    if link:
                        all_links.add(link)
                    if len(all_links) >= max_results:
                        break
                except Exception:
                    continue  
            if len(all_links) >= max_results:
                break

        driver.quit()
        return list(all_links)[:max_results]
    except Exception as e:
        st.error(f"Error searching for cases: {str(e)}")
        return []

def main():
    """Main Streamlit UI for PDF Chatbot with translation support."""
    st.set_page_config(page_title="Legal Assist", page_icon="⚖", layout="wide")
    load_css()
    
    # Get language preference
    if 'language' not in st.session_state:
        st.session_state.language = "en"
    
    # Sidebar for language selection (keep this at the top)
    with st.sidebar:
        target_lang = st.selectbox("Choose Language:", list(INDIAN_LANGUAGES.keys()), index=0)
        selected_lang_code = INDIAN_LANGUAGES[target_lang]
        st.session_state.language = selected_lang_code
        
        # Translate UI text based on selected language
        ui_text = translate_ui_text(UI_TEXT, selected_lang_code)
    
    # App title with animation
    st.markdown(
        f"<div class='app-title'>"
        f"<span class='app-icon'>⚖</span>"
        f"<h1>{translate_text(ui_text['app_title'], selected_lang_code)}</h1>"
        f"</div>", 
        unsafe_allow_html=True
    )
    
    # Initialize session state
    if 'document_text' not in st.session_state:
        st.session_state.document_text = ""

    # Sidebar continued
    with st.sidebar:
        st.markdown(f"<h2>{ui_text['settings']}</h2>", unsafe_allow_html=True)
        
        st.markdown("---")
        st.markdown(f"<h3>📄 {ui_text['document_processing']}</h3>", unsafe_allow_html=True)
        
        pdf_docs = st.file_uploader(ui_text['upload_docs'], accept_multiple_files=True, type=["pdf"])
        
        if st.button(ui_text['process_button']):
            with st.spinner(f"Processing documents..."):
                raw_text = get_pdf_text(pdf_docs)
                if not raw_text.strip():
                    st.error("No text extracted from PDFs. Please check the files.")
                else:
                    # Store the raw text for other features
                    st.session_state.document_text = raw_text
                    
                    # Create chunks with error handling
                    text_chunks = get_text_chunks(raw_text)
                    
                    # Process chunks with error handling
                    success = get_vector_store(text_chunks)
                    if success:
                        st.success(f"✅ {ui_text['processing_complete']}")
                    else:
                        st.error("There was an issue processing your documents. Please try again.")

    # Main area tabs with translated names
    tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
        f"💬 {ui_text['chat_tab']}", 
        f"📜 {ui_text['document_analysis_tab']}", 
        f"📝 {ui_text['contract_generator_tab']}", 
        f"📋 {ui_text['history_tab']}",
        f"✨ {ui_text['document_enhancement']}",
        f"🎙 {ui_text['voice_chat_tab']}",
        f"🔍 {ui_text['case_finder_tab']}"
    ])
    
    # Tab 1: Chat Interface
    with tab1:
        st.markdown("<div class='document-tab'>", unsafe_allow_html=True)
        st.markdown(f"### {ui_text['ask_question']}")
        user_question = st.text_input(ui_text['question_placeholder'])

        if user_question:
            try:
                # Get embeddings and search
                embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

                # Ensure FAISS index exists
                if not os.path.exists("faiss_index"):
                    st.error("No FAISS index found. Please upload and process PDFs first.")
                else:
                    # Load FAISS index safely
                    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
                    
                    docs = new_db.similarity_search(user_question)
                    chain = get_conversational_chain()
                    
                    response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
                    reply = response.get("output_text", "No response generated.")
                    
                    translated_reply = translate_text(reply, selected_lang_code)
                    
                    # Display in custom message container
                    st.markdown(
                        f'<div class="chat-message user"><b>{translate_text("You", selected_lang_code)}:</b> {user_question}</div>', 
                        unsafe_allow_html=True
                    )
                    st.markdown(
                        f'<div class="chat-message bot"><b>{translate_text("Legal Assistant", selected_lang_code)}:</b> {translated_reply}</div>', 
                        unsafe_allow_html=True
                    )
                    
                    # Save to history
                    save_chat_history(user_question, translated_reply)
            except Exception as e:
                st.error(f"{ui_text['error_processing']} {str(e)}")
                st.info(ui_text['try_again'])
        st.markdown("</div>", unsafe_allow_html=True)
    
    # Tab 2: Document Analysis
    with tab2:
        st.markdown("<div class='document-tab'>", unsafe_allow_html=True)
        st.markdown(f"### {ui_text['document_analysis_tab']}")
        
        if st.session_state.document_text:
            doc_text = st.session_state.document_text[:1000] + "..." if len(st.session_state.document_text) > 1000 else st.session_state.document_text
            st.text_area(ui_text['document_preview'], doc_text, height=200)
            
            analysis_options = st.selectbox(
                "Select analysis type:",
                [
                    ui_text['key_info'], 
                    ui_text['timeline'], 
                    ui_text['risk_assessment'], 
                    ui_text['legal_entities'], 
                    ui_text['obligations']
                ]
            )
            
            if st.button(ui_text['analyze_button']):
                with st.spinner("Analyzing..."):
                    try:
                        # Timeline analysis
                        if analysis_options == ui_text['timeline']:
                            timeline_data = extract_timeline(st.session_state.document_text[:50000])
                            st.markdown(f"### {ui_text['timeline_analysis']}")
                            
                            # If analysis results in a list, convert it to translated data
                            if isinstance(timeline_data, list):
                                translated_timeline = []
                                for event in timeline_data:
                                    translated_timeline.append({
                                        "date": event.get('date', 'No date'),
                                        "description": translate_text(event.get('description', 'No description'), selected_lang_code),
                                        "parties": translate_text(event.get('parties', 'Not specified'), selected_lang_code),
                                        "page": event.get('page', 'N/A')
                                    })
                                display_timeline(translated_timeline)
                            else:
                                translated_text = translate_text(timeline_data, selected_lang_code)
                                display_timeline(translated_text)
                        
                        # Legal entities analysis
                        elif analysis_options == ui_text['legal_entities']:
                            entities = extract_legal_entities(st.session_state.document_text[:50000])
                            st.markdown(f"### {ui_text['legal_entities_info']}")
                            
                            if "error" in entities:
                                st.error(translate_text(entities["error"], selected_lang_code))
                            else:
                                for category, items in entities.items():
                                    if items:  # Check if the list is not empty
                                        st.markdown(f"#### {translate_text(category.title(), selected_lang_code)}")
                                        for item in items:
                                            st.markdown(f"- {translate_text(item, selected_lang_code)}")
                        
                        # Other analyses
                        else:
                            # Map the translated option back to English for the prompt
                            analysis_prompt_map = {
                                ui_text['key_info']: "Extract key information from this legal document including: parties involved, key dates, important clauses, and main subject matter.",
                                ui_text['risk_assessment']: "Analyze this legal document for potential risks, liabilities, and areas of concern. Highlight clauses that could pose legal or financial risks.",
                                ui_text['obligations']: "Summarize the main obligations of each party in this legal document. Create a clear list of responsibilities for each identified party."
                            }
                            
                            # Use the mapped English prompt
                            analysis_prompt = analysis_prompt_map.get(analysis_options, "Extract key information from this document.")
                            
                            model = genai.GenerativeModel('gemini-1.5-flash')
                            response = model.generate_content(analysis_prompt + "\n\nDocument: " + st.session_state.document_text[:50000])
                            
                            # Translate the response
                            translated_response = translate_text(response.text, selected_lang_code)
                            
                            st.markdown(f"### {analysis_options}")
                            st.markdown(translated_response)
                    except Exception as e:
                        st.error(f"Error analyzing document: {str(e)}")
        else:
            st.info(ui_text['upload_first'])
        st.markdown("</div>", unsafe_allow_html=True)
    
    # Tab 3: Contract Generator
    with tab3:
        st.markdown("<div class='document-tab'>", unsafe_allow_html=True)
        st.markdown(f"### {ui_text['contract_generator_tab']}")
        
        contract_types = [
            ui_text['employment_agreement'], 
            ui_text['lease_agreement'], 
            ui_text['nda'], 
            ui_text['custom_contract']
        ]
        
        contract_type = st.selectbox(
            ui_text['select_contract'],
            contract_types
        )
        
        # Map selection to template names (use English keys for templates)
        contract_type_map = {
            ui_text['employment_agreement']: "employment",
            ui_text['lease_agreement']: "lease",
            ui_text['nda']: "nda",
            ui_text['custom_contract']: "custom"
        }
        
        selected_type = contract_type_map[contract_type]
        
        # Display appropriate fields based on contract type
        context_info = {"contract_type": selected_type}
        
        if selected_type == "employment":
            col1, col2 = st.columns(2)
            with col1:
                context_info["employer_name"] = st.text_input(ui_text['employer_name'])
                context_info["employer_address"] = st.text_input(ui_text['employer_address'])
                context_info["date"] = st.date_input(ui_text['agreement_date']).strftime("%B %d, %Y")
            with col2:
                context_info["employee_name"] = st.text_input(ui_text['employee_name'])
                context_info["employee_address"] = st.text_input(ui_text['employee_address'])
                context_info["position"] = st.text_input(ui_text['position_title'])
            
        elif selected_type == "lease":
            col1, col2 = st.columns(2)
            with col1:
                context_info["landlord_name"] = st.text_input(ui_text['landlord_name'])
                context_info["landlord_address"] = st.text_input(ui_text['landlord_address'])
                context_info["date"] = st.date_input(ui_text['agreement_date']).strftime("%B %d, %Y")
            with col2:
                context_info["tenant_name"] = st.text_input(ui_text['tenant_name'])
                context_info["tenant_address"] = st.text_input(ui_text['tenant_address'])
                context_info["property_details"] = st.text_area(ui_text['property_description'])
        
        elif selected_type == "nda":
            col1, col2 = st.columns(2)
            with col1:
                context_info["party1_name"] = st.text_input(ui_text['party1_name'])
                context_info["party1_address"] = st.text_input(ui_text['party1_address'])
                context_info["date"] = st.date_input(ui_text['agreement_date']).strftime("%B %d, %Y")
            with col2:
                context_info["party2_name"] = st.text_input(ui_text['party2_name'])
                context_info["party2_address"] = st.text_input(ui_text['party2_address'])
            context_info["purpose_details"] = st.text_area(ui_text['purpose_disclosure'])
        
        elif selected_type == "custom":
            context_info["description"] = st.text_area(ui_text['contract_description'])
            context_info["parties"] = st.text_input(ui_text['parties_involved'])
            context_info["date"] = st.date_input(ui_text['agreement_date']).strftime("%B %d, %Y")
        
        # Translate jurisdiction options
        jurisdiction_options = [
            translate_text("Select jurisdiction", selected_lang_code), 
            translate_text("Delhi, India", selected_lang_code), 
            translate_text("Mumbai, India", selected_lang_code), 
            translate_text("Bengaluru, India", selected_lang_code),
            translate_text("Chennai, India", selected_lang_code), 
            translate_text("Kolkata, India", selected_lang_code), 
            translate_text("Hyderabad, India", selected_lang_code)
        ]
        
        selected_jurisdiction = st.selectbox(ui_text['governing_law'], jurisdiction_options)
        
        # Map translated jurisdiction back to English for processing
        jurisdiction_map = {
            translate_text("Delhi, India", selected_lang_code): "Delhi, India",
            translate_text("Mumbai, India", selected_lang_code): "Mumbai, India",
            translate_text("Bengaluru, India", selected_lang_code): "Bengaluru, India",
            translate_text("Chennai, India", selected_lang_code): "Chennai, India",
            translate_text("Kolkata, India", selected_lang_code): "Kolkata, India",
            translate_text("Hyderabad, India", selected_lang_code): "Hyderabad, India",
        }
        
        context_info["jurisdiction"] = jurisdiction_map.get(selected_jurisdiction, selected_jurisdiction)
        
        # Place this code before the generate button to ensure preview updates on input
        if selected_type != "custom":
            # Use st.empty to create a container for the live preview that updates as you type
            preview_container = st.container()
            with preview_container:
                show_contract_preview(context_info, selected_type)
        
        if st.button(ui_text['generate_contract']) and (selected_type != "custom" or context_info.get("description")):
            with st.spinner("Generating contract..."):
                try:
                    if selected_type == "custom":
                        # Generate custom contract
                        prompt = f"""
                        Generate a professional legal contract with the following specifications:
                        - Contract description: {context_info.get('description')}
                        - Parties involved: {context_info.get('parties')}
                        - Date: {context_info.get('date')}
                        - Jurisdiction: {context_info.get('jurisdiction')}
                        
                        Create a complete, properly formatted contract with appropriate sections, clauses, and signature blocks.
                        
                        IMPORTANT: DO NOT include any markdown formatting symbols like asterisks (*), 
                        greater-than symbols (>), or bullet points. Return ONLY clean, plain text suitable 
                        for direct inclusion in a legal document. Do not include 'Option 1/2/3' language or disclaimers.
                        Format as a formal legal document with clear sections. Don't use bold or italics markers.
                        """
                        model = genai.GenerativeModel('gemini-1.5-flash')
                        response = model.generate_content(prompt)
                        
                        # Clean up the response
                        generated_contract = response.text
                        generated_contract = re.sub(r'[*>•\-]', '', generated_contract)
                        generated_contract = re.sub(r'Option \d+[:\)]', '', generated_contract)
                        generated_contract = generated_contract.replace('', '').strip()
                    else:
                        # Generate from structured template
                        generated_contract = generate_structured_contract(selected_type, context_info)
                    
                    # Translate the contract if needed (might not be desired for legal documents)
                    if selected_lang_code != "en" and st.checkbox("Translate contract to selected language"):
                        translated_contract = translate_text(generated_contract, selected_lang_code)
                    else:
                        translated_contract = generated_contract
                    
                    st.markdown(f"### {ui_text['generated_contract']}")
                    st.text_area(ui_text['preview'], translated_contract, height=300)
                    
                    # Download options
                    st.markdown("<div class='download-container'>", unsafe_allow_html=True)
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.markdown(download_document(generated_contract, f"{selected_type}_contract.txt"), unsafe_allow_html=True)
                    
                    with col2:
                        try:
                            st.markdown(download_document(generated_contract, f"{selected_type}_contract.docx"), unsafe_allow_html=True)
                        except ImportError:
                            st.info("Install python-docx for Word downloads")
                    
                    with col3:
                        try:
                            st.markdown(download_document(generated_contract, f"{selected_type}_contract.pdf"), unsafe_allow_html=True)
                        except ImportError:
                            st.info("Install reportlab for PDF downloads")
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                        
                except Exception as e:
                    st.error(f"Error generating contract: {str(e)}")
        
        st.markdown("</div>", unsafe_allow_html=True)
    
    # Tab 4: History with translated UI
    with tab4:
        st.markdown("<div class='document-tab'>", unsafe_allow_html=True)
        st.markdown(f"### {ui_text['chat_history']}")
        
        # Get and display history with translations
        history = get_chat_history()
        
        if not history:
            st.info(ui_text['no_history'])
        else:
            st.markdown("<div class='history-container'>", unsafe_allow_html=True)
            
            for item in reversed(history):  # Most recent first
                st.markdown(
                    f"<div class='history-item'>"
                    f"<div class='history-timestamp'>{item['timestamp']}</div>"
                    f"<div class='history-question'><strong>Q: {item['question']}</strong></div>"
                    f"<div class='history-answer'>A: {item['answer'][:100]}{'...' if len(item['answer']) > 100 else ''}</div>"
                    f"</div>", unsafe_allow_html=True
                )
            
            st.markdown("</div>", unsafe_allow_html=True)
        
        if st.button(ui_text['clear_history']):
            try:
                os.makedirs("history", exist_ok=True)
                with open("history/chat_history.json", "w") as f:
                    json.dump([], f)
                st.success(ui_text['history_cleared'])
                st.rerun()
            except Exception as e:
                st.error(f"Error clearing history: {str(e)}")
                
        st.markdown("</div>", unsafe_allow_html=True)
    
    # Tab 5: Document Enhancement
    with tab5:
        st.markdown("<div class='document-tab'>", unsafe_allow_html=True)
        st.markdown(f"### {ui_text['document_enhancement']}")
        
        # Document input methods
        enhancement_option = st.radio(
            ui_text['enhancement_input'],
            [ui_text['use_processed_doc'], ui_text['upload_new_doc'], ui_text['paste_text']]
        )
        
        document_to_enhance = ""
        
        if enhancement_option == ui_text['use_processed_doc']:
            if st.session_state.document_text:
                document_to_enhance = st.session_state.document_text
                st.success(ui_text['using_processed_doc'])
            else:
                st.warning(ui_text['no_processed_doc'])
        
        elif enhancement_option == ui_text['upload_new_doc']:
            uploaded_file = st.file_uploader(ui_text['upload_enhance_doc'], type=["pdf", "txt", "docx"])
            if uploaded_file is not None:
                try:
                    if uploaded_file.type == "application/pdf":
                        pdf_reader = PdfReader(uploaded_file)
                        document_to_enhance = ""
                        for page in pdf_reader.pages:
                            document_to_enhance += page.extract_text() + "\n"
                    elif uploaded_file.type == "text/plain":
                        document_to_enhance = uploaded_file.getvalue().decode("utf-8")
                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        try:
                            from docx import Document
                            doc = Document(uploaded_file)
                            document_to_enhance = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                        except ImportError:
                            st.error(ui_text['docx_import_error'])
                    
                    st.success(ui_text['document_loaded'])
                except Exception as e:
                    st.error(f"{ui_text['document_load_error']} {str(e)}")
        
        else:  # Paste text option
            document_to_enhance = st.text_area(ui_text['paste_document_text'], height=200)
        
        if document_to_enhance:
            # Preview original document
            with st.expander(ui_text['preview_original']):
                st.text_area("", document_to_enhance[:5000] + ("..." if len(document_to_enhance) > 5000 else ""), height=200)
            
            # Enhancement options
            st.markdown(f"### {ui_text['enhancement_options']}")
            
            col1, col2 = st.columns(2)
            with col1:
                apply_legal_formatting = st.checkbox(ui_text['legal_formatting'])
                apply_clause_numbering = st.checkbox(ui_text['clause_numbering'])
                apply_term_definition = st.checkbox(ui_text['term_definition'])
                apply_citation_format = st.checkbox(ui_text['citation_format'])
            
            with col2:
                apply_clarity_improvement = st.checkbox(ui_text['clarity_improvement'])
                apply_risk_identification = st.checkbox(ui_text['risk_identification'])
                apply_error_correction = st.checkbox(ui_text['error_correction'])
                apply_consistency_check = st.checkbox(ui_text['consistency_check'])
            
            # Selected enhancements summary
            enhancement_tasks = []
            if apply_legal_formatting:
                enhancement_tasks.append(ui_text['legal_formatting'])
            if apply_clause_numbering:
                enhancement_tasks.append(ui_text['clause_numbering'])
            if apply_term_definition:
                enhancement_tasks.append(ui_text['term_definition'])
            if apply_citation_format:
                enhancement_tasks.append(ui_text['citation_format'])
            if apply_clarity_improvement:
                enhancement_tasks.append(ui_text['clarity_improvement'])
            if apply_risk_identification:
                enhancement_tasks.append(ui_text['risk_identification'])
            if apply_error_correction:
                enhancement_tasks.append(ui_text['error_correction'])
            if apply_consistency_check:
                enhancement_tasks.append(ui_text['consistency_check'])
            
            if enhancement_tasks:
                st.markdown(f"{ui_text['selected_enhancements']}")
                for task in enhancement_tasks:
                    st.markdown(f"- {task}")
            else:
                st.info(ui_text['no_enhancements_selected'])
            
            # Document type (for better context)
            document_types = [
                ui_text['contract'], ui_text['legal_brief'], ui_text['court_filing'],
                ui_text['memo'], ui_text['corporate_document'], ui_text['other']
            ]
            document_type = st.selectbox(ui_text['document_type'], document_types)
            
            # Additional instructions
            additional_instructions = st.text_area(ui_text['additional_instructions'], height=100)
            
            # Enhance button
            if st.button(ui_text['enhance_document']):
                if not enhancement_tasks:
                    st.warning(ui_text['no_enhancements_selected'])
                else:
                    with st.spinner(ui_text['enhancing_document']):
                        try:
                            # Create prompt for enhancement
                            enhancement_prompt = f"""
                            You are a legal document specialist tasked with enhancing a legal document.
                            Document type: {document_type}
                            
                            Apply ONLY the following enhancements:
                            {', '.join(enhancement_tasks)}
                            
                            Additional instructions: {additional_instructions}
                            
                            Original document text:
                            {document_to_enhance[:5000000]}
                            
                            Return the enhanced document with all specified improvements.
                            Focus EXCLUSIVELY on the specified enhancements.
                            Maintain all original content while improving the document's structure and quality.
                            
                            Apply these formatting standards to the document:
                            1. Section headings should be in ALL CAPS and clearly demarcated
                            2. Proper hierarchical numbering (1, 1.1, 1.1.1, etc.)
                            3. Consistent paragraph spacing
                            4. Clear distinction between sections with line breaks
                            5. Proper formatting for defined terms (e.g., "Agreement", "Parties")
                            6. Proper indentation for lists and sub-clauses
                            7. Aligned signature blocks at the document end
                            
                            If applying legal formatting or clause numbering:
                            - Format the document with proper legal document structure
                            - Use consistent numbering convention for all clauses and sub-clauses
                            - Ensure proper indentation for hierarchical sections
                            
                            If defining legal terms:
                            - Format defined terms with initial capitals
                            - Present definitions in a consistent manner
                            - Ensure terms are properly referenced throughout
                            
                            If formatting citations:
                            - Use consistent legal citation format appropriate for the jurisdiction
                            - Format case citations, statutes, and regulations according to standard legal practice
                            
                            IMPORTANT: DO NOT summarize or add commentary about what you did.
                            ONLY return the enhanced document text, nothing else.
                            """
                            
                            # Call Gemini for document enhancement
                            model = genai.GenerativeModel('gemini-1.5-flash')
                            response = model.generate_content(enhancement_prompt)
                            
                            # Clean up the enhanced document to remove ALL markdown symbols
                            raw_text = response.text.strip()
                            enhanced_document = re.sub(r'[*>#<•\-`]', '', raw_text)  # Remove markdown symbols
                            enhanced_document = re.sub(r'\\(.?)\\*', r'\1', enhanced_document)  # Remove bold markers
                            enhanced_document = re.sub(r'(.*?)', r'\1', enhanced_document)  # Remove italic markers
                            enhanced_document = re.sub(r'.*?\n', '', enhanced_document)  # Remove code block markers
                            enhanced_document = re.sub(r'', '', enhanced_document)  # Remove closing code block markers
                            enhanced_document = re.sub(r'Option \d+[:\)]', '', enhanced_document)  # Remove option markers
                            enhanced_document = enhanced_document.strip()
                            
                            # Show enhanced document with options to download
                            st.markdown(f"### {ui_text['enhanced_document']}")
                            
                            # Display what enhancements were applied
                            st.markdown(f"{ui_text['enhancements_applied']}")
                            for task in enhancement_tasks:
                                st.markdown(f"- {task}")
                            
                            # Show the enhanced document
                            st.text_area(ui_text['preview_enhanced'], enhanced_document, height=400)
                            
                            # Download options
                            st.markdown("<div class='download-container'>", unsafe_allow_html=True)
                            col1, col2, col3 = st.columns(3)
                            
                            with col1:
                                st.markdown(download_document(enhanced_document, "enhanced_document.txt"), unsafe_allow_html=True)
                            
                            with col2:
                                try:
                                    st.markdown(download_document(enhanced_document, "enhanced_document.docx"), unsafe_allow_html=True)
                                except Exception:
                                    st.info(ui_text['docx_download_error'])
                            
                            with col3:
                                try:
                                    st.markdown(download_document(enhanced_document, "enhanced_document.pdf"), unsafe_allow_html=True)
                                except Exception:
                                    st.info(ui_text['pdf_download_error'])
                            
                            st.markdown("</div>", unsafe_allow_html=True)
                            
                            # Comparison view
                            with st.expander(ui_text['show_comparison']):
                                col1, col2 = st.columns(2)
                                with col1:
                                    st.markdown(f"{ui_text['original_document']}")
                                    st.text_area("", document_to_enhance[:5000] + 
                                             ("..." if len(document_to_enhance) > 5000 else ""), height=400)
                                
                                with col2:
                                    st.markdown(f"{ui_text['enhanced_document']}")
                                    st.text_area("", enhanced_document[:5000] + 
                                             ("..." if len(enhanced_document) > 5000 else ""), height=400)
                            
                            # Save enhanced document in session state for later use
                            st.session_state.enhanced_document = enhanced_document
                            
                        except Exception as e:
                            st.error(f"{ui_text['enhancement_error']} {str(e)}")
                            st.info(ui_text['try_again'])
        else:
            st.info(ui_text['provide_document'])
        
        st.markdown("</div>", unsafe_allow_html=True)

    # Tab 6: Voice Chat
    with tab6:
        st.markdown("<div class='document-tab'>", unsafe_allow_html=True)
        st.markdown(f"### 🎙 {ui_text['voice_chat_tab']}")
        st.markdown(f"{ui_text['voice_chat_intro']}")
        
        # Language selection for voice input/output
        voice_languages = {
            "Hindi": "hi",
            "Tamil": "ta", 
            "Telugu": "te",
            "Bengali": "bn",
            "Marathi": "mr",
            "Kannada": "kn",
            "English": "en"
        }
        
        selected_voice_lang = st.selectbox(
            ui_text['voice_language'],
            list(voice_languages.keys())
        )
        voice_lang_code = voice_languages[selected_voice_lang]
        
        # Recording duration slider
        recording_duration = st.slider(
            ui_text['recording_duration'],
            min_value=5,
            max_value=30,
            value=10,
            step=5
        )
        
        # Add environment variables info
        with st.expander("API Keys & Setup"):
            st.markdown("""
            To use the Voice Chat feature, you need these API keys:
            
            *AssemblyAI API Key* (for speech recognition)
            *ElevenLabs API Key* (for text-to-speech)
            
            Add these to your .env file:
            
            ASSEMBLYAI_API_KEY=your_assemblyai_api_key
            ELEVENLABS_API_KEY=your_elevenlabs_api_key
            
            
            Required Python packages:
            - pyaudio
            - numpy
            - soundfile
            - assemblyai
            - requests
            """)
        
        # Record button
        if st.button(ui_text['start_recording']):
            # First, check for necessary packages
            missing_packages = []
            try:
                import pyaudio
            except ImportError:
                missing_packages.append("pyaudio")
            
            try:
                import numpy as np
            except ImportError:
                missing_packages.append("numpy")
                
            try:
                import soundfile as sf
            except ImportError:
                missing_packages.append("soundfile")
                
            try:
                import assemblyai as aai
            except ImportError:
                missing_packages.append("assemblyai")
                
            if missing_packages:
                st.error(f"Missing required packages: {', '.join(missing_packages)}. Please install them to use Voice Chat.")
            else:
                # Record audio
                audio_file = record_audio(duration=recording_duration)
                
                if audio_file:
                    with st.spinner(ui_text['processing']):
                        # STT - Speech to Text
                        input_text = speech_to_text(audio_file, voice_lang_code)
                        
                        if input_text and input_text != ui_text['no_speech_detected']:
                            st.markdown(f"{ui_text['your_input']}")
                            st.write(input_text)
                            
                            # Get response using the same mechanism as text chat
                            model = genai.GenerativeModel('gemini-1.5-flash')
                            prompt = f"""
                            Respond to the following text in the same language (language code: {voice_lang_code}): 
                            {input_text}
                            
                            Provide detailed legal information in response to the query.
                            If the input is not a legal question, politely redirect and ask for a legal question.
                            Focus on laws applicable in India when appropriate.
                            Use formal but accessible language.
                            Keep your response concise but informative.
                            """
                            
                            response = model.generate_content(prompt)
                            gemini_response = response.text
                            
                            st.markdown(f"{ui_text['gemini_response']}")
                            st.write(gemini_response)
                            
                            # TTS - Text to Speech
                            output_file = text_to_speech(gemini_response, voice_lang_code)
                            
                            if not output_file:
                                output_file = fallback_text_to_speech(gemini_response, voice_lang_code)
                            
                            if output_file:
                                # Provide audio playback
                                st.audio(output_file)
                                
                                # Save chat history
                                save_chat_history(input_text, gemini_response)
                        else:
                            st.warning(ui_text['no_speech_detected'])
        
        st.markdown("</div>", unsafe_allow_html=True)
        
    # Tab 7: Legal Case Finder
    with tab7:
        st.markdown("<div class='document-tab'>", unsafe_allow_html=True)
        st.markdown(f"### 🔍 {ui_text['case_finder_title']}")
        st.markdown(f"{ui_text['case_finder_description']}")
        
        if not SELENIUM_AVAILABLE:
            st.error(ui_text['packages_missing'].format("selenium, webdriver-manager"))
            st.markdown("""
            Please install the required packages:
            
            pip install selenium webdriver-manager
            
            """)
        else:
            keyword = st.text_input(ui_text['search_keyword'])
            
            if st.button(ui_text['search_button']):
                if keyword:
                    with st.status(ui_text['searching_message'].format(keyword), expanded=True) as status:
                        links = scrape_links(keyword)
                        status.update(label="Search complete!", state="complete")
                    
                    if links:
                        st.subheader(f"📜 {ui_text['related_results']}")
                        
                        # Create a 3-column layout for results
                        cols = st.columns(3)
                        for i, link in enumerate(links):
                            domain = urllib.parse.urlparse(link).netloc  # Extract website domain
                            
                            # Styled box for each result
                            with cols[i % 3]:
                                st.markdown(
                                    f"""
                                    <div style="
                                        border: 2px solid #1a3c5a;
                                        padding: 15px;
                                        margin-bottom: 15px;
                                        border-radius: 8px;
                                        background-color: rgba(255, 255, 255, 0.05);
                                        text-align: center;
                                        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
                                        transition: transform 0.3s;
                                    ">
                                        <p><strong>{ui_text['source']}</strong> {domain}</p>
                                        <a href="{link}" target="_blank" class="download-button" style="text-decoration: none;">
                                            🔗 {ui_text['open_link']}
                                        </a>
                                    </div>
                                    """, 
                                    unsafe_allow_html=True
                                )
                    else:
                        st.warning(ui_text['no_results'])
                else:
                    st.warning(ui_text['enter_keyword'])
        
        st.markdown("</div>", unsafe_allow_html=True)

if _name_ == "_main_":
    try:
        if 'document_text' not in st.session_state:
            st.session_state.document_text = ""
        main()
    except Exception as e:
        st.error(f"Application error: {str(e)}")
        st.info("Please refresh the page and try again. If the problem persists, contact support.")